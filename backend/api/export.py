"""
导出路由
"""
from datetime import datetime, timedelta
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, and_

from backend.core.database import get_db
from backend.core.datetime_utils import get_current_time
from backend.models.recording import Recording, RecordingStatus, ScoringResult
from backend.models.rule import ScoringRule
from backend.schemas.recording import RecordingResponse
from backend.api.auth import get_current_user_required
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from urllib.parse import quote
import io

router = APIRouter(prefix="/export", tags=["导出"])


def format_seconds_to_timestamp(seconds: float) -> str:
    """将秒数转换为 mm:ss 格式的时间戳"""
    if seconds is None:
        return "00:00"
    total_seconds = int(seconds)
    minutes = total_seconds // 60
    secs = total_seconds % 60
    return f"{minutes:02d}:{secs:02d}"


def format_transcript_segments(transcript_segments, full_text=None) -> str:
    """将转写片段格式化为要求的格式：
    客服/客户【时间戳】：
    内容
    客服/客户【时间戳】：
    内容
    """
    if not transcript_segments:
        return full_text or ""

    result_lines = []
    for seg in transcript_segments:
        speaker_name = seg.get('speaker_name', seg.get('speaker', '未知'))
        start_time = seg.get('start_time', 0)
        text = seg.get('text', '')

        timestamp = format_seconds_to_timestamp(start_time)
        result_lines.append(f"{speaker_name}【{timestamp}】：\n{text}")

    return "\n".join(result_lines)


def get_date_range(days: int):
    """获取日期范围"""
    now = get_current_time()
    end = now
    start = now - timedelta(days=days - 1)
    return start.strftime('%Y-%m-%d %H:%M:%S'), end.strftime('%Y-%m-%d %H:%M:%S')


def style_cell(cell, text, bold=False, color=None, size=Pt(10)):
    """设置单元格样式"""
    cell.text = text
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = bold
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_overview_section(doc, overview_data, days_text):
    """添加概览部分"""
    doc.add_heading('数据概览', level=1)
    doc.add_paragraph(f"统计周期：{days_text}")

    # 核心指标
    doc.add_heading('核心指标', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    style_cell(table.cell(0, 0), '上传录音')
    style_cell(table.cell(0, 1), str(overview_data['total_recordings']))
    style_cell(table.cell(1, 0), '已评分')
    style_cell(table.cell(1, 1), str(overview_data['scored_count']))
    style_cell(table.cell(2, 0), '大盘平均分')
    style_cell(table.cell(2, 1), f"{overview_data['avg_total_score']}分")
    style_cell(table.cell(3, 0), '违规率')
    rate = (overview_data['recordings_with_deduction'] / overview_data['scored_count'] * 100) if overview_data['scored_count'] else 0
    style_cell(table.cell(3, 1), f"{rate:.1f}%")
    style_cell(table.cell(4, 0), '否决数')
    style_cell(table.cell(4, 1), f"{overview_data.get('recordings_with_rejection', 0)}条")

    # 加分统计
    doc.add_heading('加分统计', level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    style_cell(table2.cell(0, 0), '总加分')
    style_cell(table2.cell(0, 1), f"+{overview_data['total_bonus']}")
    style_cell(table2.cell(1, 0), '有加分的录音')
    style_cell(table2.cell(1, 1), f"{overview_data['recordings_with_bonus']}条")
    style_cell(table2.cell(2, 0), '平均加分')
    style_cell(table2.cell(2, 1), f"+{overview_data['avg_bonus']}")

    # 扣分统计
    doc.add_heading('扣分统计', level=2)
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    style_cell(table3.cell(0, 0), '总扣分')
    style_cell(table3.cell(0, 1), f"-{overview_data['total_deduction']}")
    style_cell(table3.cell(1, 0), '有扣分的录音')
    style_cell(table3.cell(1, 1), f"{overview_data['recordings_with_deduction']}条")
    style_cell(table3.cell(2, 0), '平均扣分')
    style_cell(table3.cell(2, 1), f"-{overview_data['avg_deduction']}")


def add_recordings_section(doc, recordings, title="录音详情"):
    """添加录音详情部分"""
    doc.add_heading(title, level=1)

    if not recordings:
        doc.add_paragraph('暂无数据')
        return

    # 创建表格：坐席姓名 | 文件名 | 加分 | 扣分 | 总分 | 是否否决 | 转写文本 | 扣分情况 | 加分情况
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽（单位：厘米）
    # 坐席姓名: 2cm, 文件名: 1.8cm, 加分: 0.8cm, 扣分: 0.8cm, 总分: 0.8cm, 是否否决: 1cm, 转写文本: 9cm, 扣分情况: 2.5cm, 加分情况: 2.5cm
    col_widths = [Cm(2), Cm(1.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(1), Cm(9), Cm(2.5), Cm(2.5)]

    # 表头
    headers = ['坐席姓名', '文件名', '加分', '扣分', '总分', '是否否决', '转写文本', '扣分情况', '加分情况']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.width = col_widths[i]
        style_cell(cell, header, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r in recordings:
        row = table.add_row()
        # 设置每列宽度
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

        style_cell(row.cells[0], r.get('agent_name', '-') or '-')
        style_cell(row.cells[1], r.get('file_name', '-') or '-')
        style_cell(row.cells[2], f"+{r.get('bonus_score', 0):.1f}" if r.get('bonus_score') else '+0.0')
        style_cell(row.cells[3], f"-{r.get('deduction_score', 0):.1f}" if r.get('deduction_score') else '-0.0')
        style_cell(row.cells[4], f"{r.get('total_score', 0):.1f}" if r.get('total_score') else '-')

        is_rejected = r.get('is_rejected', False)
        style_cell(row.cells[5], '是' if is_rejected else '否')

        # 转写文本（完整显示，使用格式化格式）
        formatted_transcript = format_transcript_segments(r.get('transcript_segments'), r.get('transcript'))
        style_cell(row.cells[6], formatted_transcript if formatted_transcript else '-')

        # 扣分情况 - 从scoring_details中提取，只显示命中的（分数有变化的）
        deduction_details = []
        bonus_details = []
        details = r.get('scoring_details', []) or []
        for d in details:
            detail_score = d.get('score')
            if detail_score is None:
                continue
            # 只根据分数判断，score<0是扣分，score>0是加分，0分表示未命中不显示
            item_name = d.get('item_name', d.get('rule_name', '规则'))
            if detail_score < 0:
                deduction_details.append(f"{item_name}({detail_score:.1f}分)")
            elif detail_score > 0:
                bonus_details.append(f"{item_name}(+{detail_score:.1f}分)")

        style_cell(row.cells[7], ', '.join(deduction_details) if deduction_details else '-')
        style_cell(row.cells[8], ', '.join(bonus_details) if bonus_details else '-')

    doc.add_page_break()


@router.get("/recording/{recording_id}")
async def export_single_recording_report(
    recording_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user_required)
):
    """导出单条录音的报告"""
    user_id = current_user.get("loginid", "admin")

    # 获取录音
    result = await db.execute(
        select(Recording).where(
            Recording.id == recording_id,
            Recording.user_id == user_id
        )
    )
    recording = result.scalars().first()
    if not recording:
        return {"error": "录音不存在"}

    # 获取评分详情
    score_result = await db.execute(
        select(ScoringResult).where(ScoringResult.recording_id == recording_id)
    )
    scoring = score_result.scalars().first()

    doc = Document()

    # 标题
    doc.add_heading('录音评分报告', level=1)

    # 基本信息
    doc.add_heading('基本信息', level=2)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    style_cell(info_table.cell(0, 0), '文件名')
    style_cell(info_table.cell(0, 1), recording.file_name or '-')
    style_cell(info_table.cell(1, 0), '坐席姓名')
    style_cell(info_table.cell(1, 1), recording.agent_name or '-')
    style_cell(info_table.cell(2, 0), '坐席工号')
    style_cell(info_table.cell(2, 1), recording.agent_id or '-')
    style_cell(info_table.cell(3, 0), '通话时间')
    style_cell(info_table.cell(3, 1), recording.call_time.strftime('%Y-%m-%d %H:%M:%S') if recording.call_time else '-')
    style_cell(info_table.cell(4, 0), '录音时长')
    style_cell(info_table.cell(4, 1), f"{recording.duration or 0}秒" if recording.duration else '-')
    style_cell(info_table.cell(5, 0), '总分')
    score_val = recording.total_score if recording.total_score is not None else 0
    style_cell(info_table.cell(5, 1), f"{score_val}分")

    # 评分结果
    doc.add_heading('评分结果', level=2)
    score_table = doc.add_table(rows=3, cols=2)
    score_table.style = 'Table Grid'
    bonus = recording.bonus_score or 0
    deduction = recording.deduction_score or 0
    style_cell(score_table.cell(0, 0), '加分')
    style_cell(score_table.cell(0, 1), f"+{bonus}")
    style_cell(score_table.cell(1, 0), '扣分')
    style_cell(score_table.cell(1, 1), f"-{deduction}")
    style_cell(score_table.cell(2, 0), '总分')
    style_cell(score_table.cell(2, 1), f"{score_val}")

    # 评分明细
    if scoring and scoring.scoring_details:
        doc.add_heading('评分明细', level=2)
        detail_table = doc.add_table(rows=1, cols=5)
        detail_table.style = 'Table Grid'
        headers = ['考核项', '类型', '状态', '得分', '匹配文本']
        header_row = detail_table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            style_cell(cell, header, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for d in scoring.scoring_details:
            row = detail_table.add_row()
            item_type = d.get('item_type', 'bonus')
            status = d.get('status', 'not_matched')
            score = d.get('score', 0)
            max_score = d.get('max_score', 0)

            style_cell(row.cells[0], d.get('item_name', '-'))
            style_cell(row.cells[1], '加分' if item_type == 'bonus' else '扣分')
            status_map = {'matched': '已匹配', 'not_matched': '未匹配'}
            style_cell(row.cells[2], status_map.get(status, status))
            style_cell(row.cells[3], f"{score}/{max_score}")
            style_cell(row.cells[4], d.get('matched_text', '-'))

    # 转写文本
    if recording.transcript or recording.transcript_segments:
        doc.add_heading('转写文本', level=2)
        formatted_transcript = format_transcript_segments(recording.transcript_segments, recording.transcript)
        doc.add_paragraph(formatted_transcript)

    # 保存文档
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{get_current_time().strftime('%Y%m%d')}-{recording.agent_name or '未知'}-{recording.file_name}.docx"

    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{quote(filename)}"'}
    )


@router.get("/report")
async def export_report(
    type: str = Query("all", description="all:整体导出, agent:坐席导出"),
    agent_name: str = Query(None, description="坐席姓名（当type=agent时必传）"),
    start_date: str = Query(None, description="开始日期"),
    end_date: str = Query(None, description="结束日期"),
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user_required)
):
    """导出评分报告（按用户隔离）"""
    user_id = current_user.get("loginid", "admin")
    doc = Document()

    # 日期范围
    date_ranges = []
    if start_date and end_date:
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d %H:%M:%S')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d %H:%M:%S')
            date_ranges = [(start_dt, end_dt, f"{start_date[:10]} to {end_date[:10]}")]
        except:
            date_ranges = [
                (get_current_time() - timedelta(days=6), get_current_time(), 'Last 7 days'),
                (get_current_time() - timedelta(days=29), get_current_time(), 'Last 30 days')
            ]
    else:
        date_ranges = [
            (get_current_time() - timedelta(days=6), get_current_time(), 'Last 7 days'),
            (get_current_time() - timedelta(days=29), get_current_time(), 'Last 30 days')
        ]

    for start_dt, end_dt, days_text in date_ranges:
        conditions = [
            Recording.status == RecordingStatus.SCORED,
            Recording.total_score.isnot(None),
            Recording.created_at >= start_dt,
            Recording.created_at < end_dt + timedelta(days=1),
            Recording.user_id == user_id
        ]

        if type == "agent" and agent_name:
            conditions.append(Recording.agent_name == agent_name)

        # 获取概览数据 - total_recordings（所有上传的录音，不限状态）
        total_conditions = [
            Recording.user_id == user_id,
            Recording.created_at >= start_dt,
            Recording.created_at < end_dt + timedelta(days=1)
        ]
        if type == "agent" and agent_name:
            total_conditions.append(Recording.agent_name == agent_name)
        total_result = await db.execute(select(func.count()).where(and_(*total_conditions)))
        total_recordings = total_result.scalar() or 0

        # 已评分数
        scored_count = await db.execute(select(func.count()).where(and_(*conditions)))
        scored_count = scored_count.scalar() or 0

        bonus_cond = conditions + [Recording.bonus_score.isnot(None), Recording.bonus_score > 0]
        bonus_count_result = await db.execute(select(func.count()).where(and_(*bonus_cond)))
        recordings_with_bonus = bonus_count_result.scalar() or 0

        deduct_cond = conditions + [Recording.deduction_score.isnot(None), Recording.deduction_score > 0]
        deduct_count_result = await db.execute(select(func.count()).where(and_(*deduct_cond)))
        recordings_with_deduction = deduct_count_result.scalar() or 0

        total_bonus_result = await db.execute(select(func.coalesce(func.sum(Recording.bonus_score), 0)).where(and_(*conditions)))
        total_bonus = total_bonus_result.scalar() or 0

        total_deduct_result = await db.execute(select(func.coalesce(func.sum(Recording.deduction_score), 0)).where(and_(*conditions)))
        total_deduction = total_deduct_result.scalar() or 0

        avg_bonus_result = await db.execute(select(func.avg(Recording.bonus_score)).where(and_(*bonus_cond)))
        avg_bonus = avg_bonus_result.scalar() or 0

        avg_deduct_result = await db.execute(select(func.avg(Recording.deduction_score)).where(and_(*deduct_cond)))
        avg_deduction = avg_deduct_result.scalar() or 0

        avg_total_result = await db.execute(select(func.avg(Recording.total_score)).where(and_(*conditions)))
        avg_total = avg_total_result.scalar() or 0

        reject_cond = conditions + [ScoringResult.is_rejected == True]
        reject_count_result = await db.execute(
            select(func.count())
            .select_from(Recording)
            .join(ScoringResult, ScoringResult.recording_id == Recording.id)
            .where(and_(*reject_cond))
        )
        recordings_with_rejection = reject_count_result.scalar() or 0

        overview_data = {
            'total_recordings': total_recordings,
            'scored_count': scored_count,
            'recordings_with_bonus': recordings_with_bonus,
            'recordings_with_deduction': recordings_with_deduction,
            'recordings_with_rejection': recordings_with_rejection,
            'total_bonus': round(total_bonus, 1),
            'total_deduction': round(total_deduction, 1),
            'avg_bonus': round(avg_bonus, 1) if avg_bonus else 0,
            'avg_deduction': round(avg_deduction, 1) if avg_deduction else 0,
            'avg_total_score': round(avg_total, 1) if avg_total else 0,
        }

        # 添加概览部分
        add_overview_section(doc, overview_data, days_text)

        # 获取录音列表（分页获取所有）
        page = 1
        page_size = 100
        all_recordings = []

        while True:
            query = select(Recording).where(and_(*conditions)).order_by(Recording.created_at.desc()).offset((page - 1) * page_size).limit(page_size)
            result = await db.execute(query)
            recordings = result.scalars().all()

            if not recordings:
                break

            for recording in recordings:
                # 获取评分详情
                score_result = await db.execute(select(ScoringResult).where(ScoringResult.recording_id == recording.id))
                scoring = score_result.scalars().first()

                all_recordings.append({
                    'agent_name': recording.agent_name or '未知',
                    'file_name': recording.file_name or '-',
                    'bonus_score': recording.bonus_score or 0,
                    'deduction_score': recording.deduction_score or 0,
                    'total_score': recording.total_score or 0,
                    'transcript': recording.transcript or '',
                    'transcript_segments': recording.transcript_segments or [],
                    'scoring_details': scoring.scoring_details if scoring and scoring.scoring_details else [],
                    'is_rejected': getattr(scoring, 'is_rejected', False) if scoring else False,
                    'created_at': recording.created_at,
                })

            if len(recordings) < page_size:
                break
            page += 1

        # 添加录音详情部分
        title = f"Recording Details{' - ' + agent_name if agent_name else ''} ({days_text})"
        add_recordings_section(doc, all_recordings, title)

    # 保存文档
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"score_report_{agent_name if agent_name else 'all'}_{get_current_time().strftime('%Y%m%d%H%M%S')}.docx"

    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{quote(filename)}"'}
    )


@router.get("/agents")
async def get_agent_list(
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user_required)
):
    """获取坐席列表（用于导出时的选择，按用户隔离）"""
    user_id = current_user.get("loginid", "admin")
    result = await db.execute(
        select(Recording.agent_name)
        .where(
            Recording.agent_name.isnot(None),
            Recording.user_id == user_id
        )
        .group_by(Recording.agent_name)
        .order_by(Recording.agent_name)
    )
    agents = result.scalars().all()
    return [{'agent_name': name} for name in agents if name]